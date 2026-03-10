import re
from docx import Document
from datetime import date

padrao_log = re.compile(r"^(\d{4}-\d{2}-\d{2})\s(\d{2}:\d{2}:\d{2})\s(INFO|ERROR|WARNING|DEBUG)\s+(.*)$")

contagem_total = {"INFO": 0, "ERROR": 0, "WARNING": 0, "DEBUG": 0}
mensagens_erro = []
contagem_por_dia = {} 

print("Analisando o arquivo de log...")

try:
    with open("logs.txt", "r", encoding="utf-8") as arquivo_txt:
        for linha in arquivo_txt:
            match = padrao_log.match(linha.strip())
            
            if match:
                data_log = match.group(1)
                hora_log = match.group(2)
                nivel = match.group(3)
                mensagem = match.group(4)
                
                contagem_total[nivel] += 1
                if nivel == "ERROR":
                    mensagens_erro.append(f"[{data_log} {hora_log}] {mensagem}")

                if data_log not in contagem_por_dia:
                    contagem_por_dia[data_log] = {"INFO": 0, "ERROR": 0, "WARNING": 0, "DEBUG": 0}
                
                contagem_por_dia[data_log][nivel] += 1

except FileNotFoundError:
    print("Erro: O arquivo 'server.log' não foi encontrado na pasta.")
    exit()

print("Gerando o relatório Word...")
doc = Document()

doc.add_heading("Relatório de Análise de Logs", level=1)

resumo = (
    f"A análise dos logs foi concluída com sucesso. No total, foram registrados: "
    f"{contagem_total['INFO']} logs de INFO, "
    f"{contagem_total['WARNING']} de WARNING, "
    f"{contagem_total['DEBUG']} de DEBUG e "
    f"{contagem_total['ERROR']} de ERROR."
)
doc.add_paragraph(resumo)

doc.add_heading("Detalhamento de Erros", level=2)
if mensagens_erro:
    for erro in mensagens_erro:
        doc.add_paragraph(erro, style="List Bullet")
else:
    doc.add_paragraph("Nenhum erro foi registrado no período analisado.")

doc.add_paragraph()

doc.add_heading("Registros por Dia", level=2)

tabela = doc.add_table(rows=1, cols=5)
tabela.style = "Table Grid"

cabecalho = tabela.rows[0].cells
cabecalho[0].text = "Data"
cabecalho[1].text = "INFO"
cabecalho[2].text = "WARNING"
cabecalho[3].text = "ERROR"
cabecalho[4].text = "DEBUG"

for data, niveis in contagem_por_dia.items():
    linha_celulas = tabela.add_row().cells
    linha_celulas[0].text = data
    linha_celulas[1].text = str(niveis["INFO"])
    linha_celulas[2].text = str(niveis["WARNING"])
    linha_celulas[3].text = str(niveis["ERROR"])
    linha_celulas[4].text = str(niveis["DEBUG"])
data_hoje = date.today().strftime("%Y-%m-%d")
nome_arquivo = f"relatorio_logs_{data_hoje}.docx"

doc.save(nome_arquivo)
print(f"Sucesso! Relatório salvo como '{nome_arquivo}'.")